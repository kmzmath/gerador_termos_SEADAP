# ───────────────────────── IMPORTS ───────────────────────── #
from decimal import Decimal, InvalidOperation, getcontext
from pathlib import Path
import io
import locale
import re
from typing import Dict

import streamlit as st
from dateutil.parser import parse, ParserError
from docx import Document
from num2words import num2words

# ───────────────────── CONFIGURAÇÃO GLOBAL ───────────────────── #
getcontext().prec = 12
try:
    locale.setlocale(locale.LC_TIME, "pt_BR.UTF-8")
except locale.Error:
    pass

TEMPLATE_FILE = Path("Template_RECUPERA EXPRESSv2.docx")

# ───────────────────── FUNÇÕES DE FORMATO ───────────────────── #
def br_currency(value: str) -> str:
    try:
        num = Decimal(value.replace(".", "").replace(",", ".")).quantize(Decimal("0.01"))
        return f"{num:,.2f}".replace(",", "TEMP").replace(".", ",").replace("TEMP", ".")
    except (InvalidOperation, AttributeError):
        return value


def format_cnpj(cnpj: str) -> str:
    d = re.sub(r"\D", "", str(cnpj))
    return f"{d[:2]}.{d[2:5]}.{d[5:8]}/{d[8:12]}-{d[12:]}" if len(d) == 14 else cnpj


def format_cpf(cpf: str) -> str:
    d = re.sub(r"\D", "", str(cpf))
    return f"{d[:3]}.{d[3:6]}.{d[6:9]}-{d[9:]}" if len(d) == 11 else cpf


def date_ddmmyyyy(d) -> str:
    try:
        return parse(str(d), dayfirst=True).strftime("%d/%m/%Y")
    except (ParserError, ValueError, TypeError):
        return str(d)


def month_year(d) -> str:
    try:
        return parse(str(d), dayfirst=True).strftime("%B de %Y").capitalize()
    except (ParserError, ValueError, TypeError):
        return str(d)


def numero_para_extenso_completo(num, modo="uif") -> str:
    try:
        inteiro_str, dec_str = br_currency(str(num)).split(",")
        inteiro, dec = int(inteiro_str.replace(".", "")), int(dec_str)
        ext_i = num2words(inteiro, lang="pt_BR")
        if dec:
            ext_d = num2words(dec, lang="pt_BR")
            if modo == "uif":
                suf_i = "inteiro" if inteiro == 1 else "inteiros"
                suf_d = "centésimo" if dec == 1 else "centésimos"
                return f"{ext_i} {suf_i}, e {ext_d} {suf_d}"
            return f"{ext_i} vírgula {ext_d}"
        return ext_i
    except Exception:
        return str(num)

# ──────── SUBSTITUIÇÃO QUE PRESERVA FORMATAÇÃO ───────── #
def docx_replace(doc: Document, mapping: Dict[str, str]) -> Document:
    if not mapping:
        return doc
    for para in doc.paragraphs:
        _replace_in_para(para, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para, mapping)
    return doc


def _replace_in_para(paragraph, mapping: Dict[str, str]):
    for placeholder, new_val in mapping.items():
        if placeholder not in paragraph.text:
            continue
        while placeholder in paragraph.text:
            runs = paragraph.runs
            full = "".join(r.text for r in runs)
            start = full.find(placeholder)
            end = start + len(placeholder)

            cur = 0
            run_start = run_end = 0
            for idx, run in enumerate(runs):
                nxt = cur + len(run.text)
                if start >= cur and start < nxt:
                    run_start = idx
                if end > cur and end <= nxt:
                    run_end = idx
                    break
                cur = nxt

            off_start = start - sum(len(r.text) for r in runs[:run_start])
            off_end = end - sum(len(r.text) for r in runs[:run_end])

            if run_start == run_end:
                r = runs[run_start]
                r.text = r.text[:off_start] + str(new_val) + r.text[off_end:]
            else:
                runs[run_start].text = runs[run_start].text[:off_start] + str(new_val)
                for i in range(run_start + 1, run_end):
                    runs[i].text = ""
                runs[run_end].text = runs[run_end].text[off_end:]

# ─────────────────── HELPERS DE PLACEHOLDER ─────────────────── #
def add_valor(orig: str, valor: str, rep: Dict[str, str], ref: str | None = None):
    if not valor:
        return
    fmt = br_currency(valor)
    ext = numero_para_extenso_completo(fmt)
    rep[orig] = f"{fmt} UIF/RS ({ext} de Unidades de Incentivo do FUNDOPEM/RS)"
    if ref:
        rep[ref] = fmt

# ───────────────────── INTERFACE STREAMLIT ───────────────────── #
st.set_page_config(page_title="Gerador Termos FUNDOPEM", layout="wide", page_icon="📄")

st.markdown(
    """
    <style>
      input:focus, select:focus, textarea:focus {
        outline: 3px solid #ff922b !important;
        border: 2px solid #ff922b !important;
      }
    </style>
    """,
    unsafe_allow_html=True,
)

st.title("📄 Gerador Automatizado de Termos de Ajuste")

if not TEMPLATE_FILE.exists():
    st.error(f"Template não encontrado: {TEMPLATE_FILE.resolve()}")
    st.stop()

# -------- Formulário -------- #
with st.form("dados_termo"):
    st.subheader("1. Informações Gerais")
    c1, c2 = st.columns(2)
    with c1:
        termo_num           = st.text_input("Número do termo")
        empresa_nome        = st.text_input("Nome da Empresa")
        empresa_cnpj        = st.text_input("CNPJ")
        empresa_cgcte       = st.text_input("CGC/TE")
        empresa_endereco    = st.text_input("Endereço Completo")
        proa_num            = st.text_input("Nº do Processo (PROA)")
    with c2:
        representante_nome  = st.text_input("Nome do Representante Legal")
        representante_cpf   = st.text_input("CPF do Representante")
        parecer_num         = st.text_input("Nº do Parecer")
        parecer_data        = st.date_input("Data do Parecer", format="DD/MM/YYYY")
        doe_data            = st.date_input("Data do DOE", format="DD/MM/YYYY")

    st.subheader("2. Detalhes do Projeto")
    c3, c4 = st.columns(2)
    with c3:
        empresa_porte   = st.selectbox("Porte da Empresa", ("Pequeno", "Médio", "Grande"))
        municipio       = st.text_input("Município")
        corede          = st.text_input("COREDE")
        qtd_empregos    = st.text_input("Quantidade de Empregos", placeholder="1.000")
    with c4:
        pontos_fundopem     = st.text_input("Pontos FUNDOPEM (3.1)")
        set_estrategicos    = st.text_input("Pontos Set. Estratégicos")
        intensidade_tec     = st.text_input("Pontos Intensidade Tecnológica")
        perc_integrar       = st.text_input("Percentual INTEGRAR (3.2)", placeholder="34,55")
        pontos_idese        = st.text_input("Pontos IDESE")
        pontos_setor        = st.text_input("Pontos Setor Industrial")

    st.subheader("3. Valores e Prazos")
    c5, c6 = st.columns(2)
    with c5:
        valor_total             = st.text_input("Valor Total do Projeto (2.1)")
        valor_apres_inicial     = st.text_input("Valor Apresentado Inicialmente (2.3)")
        valor_inicial_aceito    = st.text_input("Valor Inicialmente Aceito (2.3.1)")
        equips_24               = st.text_input("Equipamentos (2.4)")
        limite_max_liberado     = st.text_input("Limite Máximo Liberado (4.1.2)")
        valor_liberado_fruicao  = st.text_input("Valor Liberado p/ Fruição (4.1.2.1)")
    with c6:
        data_inicio         = st.date_input("Início da Vigência", format="DD/MM/YYYY")
        data_final_fruicao  = st.date_input("Final da Fruição", format="DD/MM/YYYY")
        mes_regularidade    = st.text_input("Mês da Regularidade (ex: Julho/2025)")

    # ---- Botões lado a lado ---- #
    bcol1, bcol2 = st.columns([2, 1])
    with bcol1:
        gerar = st.form_submit_button("🚀 Gerar Documento Word", type="primary")
    with bcol2:
        limpar = st.form_submit_button("🗑️ Limpar formulário")

# -------- Ações fora do form -------- #
if limpar:
    st.session_state.clear()
    (st.rerun if hasattr(st, "rerun") else st.experimental_rerun)()

if gerar:
    with st.spinner("Gerando documento..."):
        rep: Dict[str, str] = {}

        # ---------------- Valores UIF/RS ---------------- #
        add_valor(
            "693.224,32 UIF/RS (seiscentos e noventa e três mil, duzentos e vinte e quatro inteiros, e trinta e dois centésimos de Unidades de Incentivo do FUNDOPEM/RS)",
            valor_total, rep
        )
        add_valor(
            "193.874,41 UIF/RS (cento e noventa e três mil, oitocentos e setenta e quatro inteiros, e quarenta e um centésimos de Unidades de Incentivo do FUNDOPEM/RS)",
            valor_apres_inicial, rep, ref="=g10"
        )
        add_valor(
            "159.123,22 UIF/RS (cento e cinquenta e nove mil, cento e vinte e três inteiros, e vinte e dois centésimos de Unidades de Incentivo do FUNDOPEM/RS)",
            valor_inicial_aceito, rep, ref="=g11"
        )
        if equips_24:
            rep[
                "Do valor estabelecido no item 2.3.1 desta Cláusula, o montante de 193.874,41 UIF/RS (cento e noventa e três mil, oitocentos e setenta e quatro inteiros, e quarenta e um centésimos de Unidades de Incentivo do FUNDOPEM/RS) contempla os investimentos realizados em equipamentos."
            ] = (
                f"Do valor estabelecido no item 2.3.1 desta Cláusula, o montante de {br_currency(equips_24)} UIF/RS "
                f"({numero_para_extenso_completo(br_currency(equips_24))} de Unidades de Incentivo do FUNDOPEM/RS) contempla os investimentos realizados em equipamentos."
            )
        add_valor(
            "239.509,00 UIF/RS (duzentos e trinta e nove mil, e quinhentos e nove inteiros de Unidades de Incentivo do FUNDOPEM/RS)",
            limite_max_liberado, rep, ref="=g8"
        )
        add_valor(
            "62.299,92 UIF/RS (sessenta e dois mil, duzentos e noventa e nove inteiros, e noventa e dois centésimos de Unidades de Incentivo do FUNDOPEM/RS)",
            valor_liberado_fruicao, rep, ref="=g21"
        )

        # ---------------- Pontos / Percentual ------------- #
        if pontos_fundopem:
            rep["#pontos# (sessenta) pontos"] = (
                f"{pontos_fundopem} ({numero_para_extenso_completo(pontos_fundopem, 'geral')}) pontos"
            )
        if perc_integrar:
            perc_ext = numero_para_extenso_completo(perc_integrar, "percent") + " por cento"
            rep["#integrar# (trinta e quatro vírgula cinquenta e cinco por cento)"] = f"{perc_integrar}% ({perc_ext})"
            rep["#integrar#%"] = f"{perc_integrar}%"
            rep["#integrar#"] = perc_integrar

        # ---------------- Placeholders simples ------------- #
        rep.update(
            {
                "#xx/aaaa#": termo_num,
                "#EMPRESA#": empresa_nome.upper(),
                "#ENDERECO#": empresa_endereco,
                "#XX.XXX.XXX/0001-XX#": format_cnpj(empresa_cnpj),
                "#REPRESENTANTE#": representante_nome,
                "#cpf#": format_cpf(representante_cpf),
                "#proa#": proa_num,
                "#MUNICIPIO_E_COREDE#": f"{municipio}/RS (COREDE: {corede})" if municipio and corede else "",
                "#idese#": pontos_idese,
                "#setint#": pontos_setor,
                "#set#": set_estrategicos,
                "#it#": intensidade_tec,
                "#porte#": empresa_porte,
                "#cgcte#": empresa_cgcte,
                "#pontos#": pontos_fundopem,
            }
        )

        # ---------------- Datas ---------------------------- #
        if data_inicio:
            rep["#inicio#"] = date_ddmmyyyy(data_inicio)
        if data_final_fruicao:
            rep["#final#"] = month_year(data_final_fruicao)
            rep["#final2#"] = month_year(data_final_fruicao).replace(" de ", "/")
        if mes_regularidade:
            rep["#regularidade#"] = mes_regularidade
        if parecer_num and parecer_data and doe_data:
            rep["Parecer nº xxx/aaaa, de dd.mm.aaaa (DOE de dd.mm.aaaa)"] = (
                f"Parecer Nº {parecer_num}, de {date_ddmmyyyy(parecer_data)} "
                f"(DOE de {date_ddmmyyyy(doe_data)})"
            )
        if qtd_empregos:
            rep["#emp#"] = str(qtd_empregos)
            rep["(duzentos e noventa e sete)"] = (
                f"({numero_para_extenso_completo(qtd_empregos, 'geral')})"
            )

        # ---------------- Geração do DOCX ------------------ #
        doc = Document(str(TEMPLATE_FILE))
        docx_replace(doc, rep)

        buf = io.BytesIO()
        doc.save(buf)

        nome_arquivo = (
            f"Termo_Ajuste_{(termo_num or 'novo').replace('/', '-')}_"
            f"{(empresa_nome or 'empresa').replace(' ', '_')}.docx"
        )
        st.success("🎉 Documento gerado com sucesso!")
        st.download_button(
            "⬇️ Baixar documento Word",
            data=buf.getvalue(),
            file_name=nome_arquivo,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )