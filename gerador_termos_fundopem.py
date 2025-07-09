# ─────────────────── IMPORTS ─────────────────── #
from decimal import Decimal, InvalidOperation, getcontext
from pathlib import Path
import io, locale, re
from typing import Dict, Optional
from datetime import date, datetime

import streamlit as st
from dateutil.parser import parse, ParserError
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from num2words import num2words

# ───────────────── GLOBALS ───────────────── #
getcontext().prec = 12
try:
    locale.setlocale(locale.LC_TIME, "pt_BR.UTF-8")
except locale.Error:
    pass

TEMPLATE_FILE = Path("TEMPLATE_RECUPERA EXPRESS.docx")

# ───────────────── UTILIDADES ───────────────── #
def br_currency(v: str) -> str:
    try:
        num = Decimal(v.replace(".", "").replace(",", ".")).quantize(Decimal("0.01"))
        return f"{num:,.2f}".replace(",", "TEMP").replace(".", ",").replace("TEMP", ".")
    except (InvalidOperation, AttributeError):
        return v


def format_cnpj(cnpj: str) -> str:
    d = re.sub(r"\D", "", str(cnpj))
    return f"{d[:2]}.{d[2:5]}.{d[5:8]}/{d[8:12]}-{d[12:]}" if len(d) == 14 else cnpj


def format_cpf(cpf: str) -> str:
    d = re.sub(r"\D", "", str(cpf))
    return f"{d[:3]}.{d[3:6]}.{d[6:9]}-{d[9:]}" if len(d) == 11 else cpf


# Datas em pt-BR
MESES = [
    "janeiro", "fevereiro", "março", "abril", "maio", "junho",
    "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"
]

def to_date(x) -> Optional[date]:
    if isinstance(x, date) and not isinstance(x, datetime):
        return x
    if isinstance(x, datetime):
        return x.date()
    if isinstance(x, str):
        # Tenta diferentes formatos
        for fmt in ["%d/%m/%Y", "%d.%m.%Y", "%Y-%m-%d"]:
            try:
                return datetime.strptime(x, fmt).date()
            except ValueError:
                continue
    try:
        return parse(str(x), dayfirst=True).date()
    except (ParserError, ValueError):
        return None

def data_texto(dt: date, mes_ano=False) -> str:
    if dt is None:
        return ""
    if mes_ano:
        return f"{MESES[dt.month-1].capitalize()}/{dt.year}"
    return f"{dt.day} de {MESES[dt.month-1]} de {dt.year}"

def data_formato_ponto(dt: date) -> str:
    """Formata data no formato DD.MM.AAAA"""
    if dt is None:
        return ""
    return f"{dt.day:02d}.{dt.month:02d}.{dt.year}"


def num_extenso(num, modo="uif"):
    try:
        int_str, dec_str = br_currency(str(num)).split(",")
        inteiro, dec = int(int_str.replace(".", "")), int(dec_str)
        ext_i = num2words(inteiro, lang="pt_BR")
        if dec:
            ext_d = num2words(dec, lang="pt_BR")
            if modo == "uif":
                return f"{ext_i} {'inteiro' if inteiro==1 else 'inteiros'}, e {ext_d} {'centésimo' if dec==1 else 'centésimos'}"
            return f"{ext_i} vírgula {ext_d}"
        return ext_i
    except Exception:
        return str(num)


def docx_replace(doc: Document, mapping: Dict[str, str]) -> Document:
    """
    Substitui placeholders em um documento .docx, preservando a formatação
    mesmo quando o placeholder se estende por múltiplos 'runs' com estilos diferentes.
    """
    if not mapping:
        return doc

    def replace_in_paragraph(paragraph: Paragraph):
        for placeholder, value in mapping.items():
            if placeholder not in paragraph.text:
                continue
            
            # Repete enquanto o placeholder for encontrado, para múltiplas ocorrências
            while placeholder in paragraph.text:
                runs = paragraph.runs
                full_text = "".join(r.text for r in runs)
                
                # Encontra a posição do placeholder no texto completo do parágrafo
                start_pos = full_text.find(placeholder)
                end_pos = start_pos + len(placeholder)

                # Encontra os índices dos runs de início e fim
                current_pos = 0
                run_start_idx, run_end_idx = -1, -1
                for i, run in enumerate(runs):
                    next_pos = current_pos + len(run.text)
                    if run_start_idx == -1 and start_pos >= current_pos and start_pos < next_pos:
                        run_start_idx = i
                    if run_end_idx == -1 and end_pos > current_pos and end_pos <= next_pos:
                        run_end_idx = i
                        break
                    current_pos = next_pos
                
                # Calcula o offset dentro do run de início e fim
                offset_start = start_pos - sum(len(r.text) for r in runs[:run_start_idx])
                offset_end = end_pos - sum(len(r.text) for r in runs[:run_end_idx])

                # Caso 1: Placeholder contido em um único run
                if run_start_idx == run_end_idx:
                    run = runs[run_start_idx]
                    run.text = run.text[:offset_start] + str(value) + run.text[offset_end:]
                
                # Caso 2: Placeholder se estende por múltiplos runs
                else:
                    text_to_distribute = str(value)
                    
                    # 1. Trata o primeiro run
                    first_run = runs[run_start_idx]
                    space_in_first_run = len(first_run.text) - offset_start
                    chunk = text_to_distribute[:space_in_first_run]
                    first_run.text = first_run.text[:offset_start] + chunk
                    text_to_distribute = text_to_distribute[space_in_first_run:]

                    # 2. Trata os runs intermediários
                    for i in range(run_start_idx + 1, run_end_idx):
                        intermediate_run = runs[i]
                        space_in_intermediate_run = len(intermediate_run.text)
                        chunk = text_to_distribute[:space_in_intermediate_run]
                        intermediate_run.text = chunk
                        text_to_distribute = text_to_distribute[space_in_intermediate_run:]

                    # 3. Trata o último run
                    last_run = runs[run_end_idx]
                    tail = last_run.text[offset_end:]
                    last_run.text = text_to_distribute + tail

    # Itera por todos os parágrafos no corpo do documento e nas tabelas
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    return doc
# ==============================================================================

def add_valor(orig: str, v: str, mp: Dict[str, str], ref: str | None = None):
    if v:
        fmt = br_currency(v)
        mp[orig] = f"{fmt} UIF/RS ({num_extenso(fmt)} de Unidades de Incentivo do FUNDOPEM/RS)"
        if ref:
            mp[ref] = fmt

# ─────────── INTERFACE ─────────── #
st.set_page_config(page_title="Gerador Termos FUNDOPEM",
                   page_icon="📄",
                   layout="wide")

st.markdown(
    "<style>input:focus,select:focus,textarea:focus{outline:3px solid #ff922b!important;border:2px solid #ff922b!important}</style>",
    unsafe_allow_html=True)

# Versão e título
col_title, col_version = st.columns([4, 1])
with col_title:
    st.title("📄 Gerador Automatizado de Termos de Ajuste")
with col_version:
    st.markdown(
        """
        <div style='text-align: right; padding-top: 20px;'>
            <span style='background-color: #ff922b; color: white; padding: 5px 10px; border-radius: 5px; font-size: 14px; font-weight: bold;'>
                v2.0.1
            </span>
            <br>
            <span style='font-size: 12px; color: #666;'>09/07/2025</span>
        </div>
        """,
        unsafe_allow_html=True
    )

if not TEMPLATE_FILE.exists():
    st.error(f"Template não encontrado: {TEMPLATE_FILE.resolve()}")
    st.stop()

# ─────────── FORMULÁRIO ─────────── #
with st.form("dados"):
    st.subheader("1. Informações Gerais")
    c1, c2 = st.columns(2)
    with c1:
        termo_num = st.text_input("Número do termo")
        empresa_nome = st.text_input("Nome da Empresa")
        empresa_cnpj = st.text_input("CNPJ")
        empresa_cgcte = st.text_input("CGC/TE")
        empresa_endereco = st.text_input("Endereço Completo")
        proa_num = st.text_input("Nº do Processo (PROA)")
        proa_data = st.text_input("Data do Processo (DD/MM/AAAA)", placeholder="26/12/2024")
    with c2:
        representante_nome = st.text_input("Nome do Representante Legal")
        representante_cpf = st.text_input("CPF do Representante")
        parecer_num = st.text_input("Nº do Parecer")
        parecer_data = st.text_input("Data do Parecer (DD/MM/AAAA)", placeholder="24/04/2025")
        doe_data = st.text_input("Data do DOE (DD/MM/AAAA)", placeholder="05/05/2025")

    st.subheader("2. Detalhes do Projeto")
    c3, c4 = st.columns(2)
    with c3:
        empresa_porte = st.selectbox("Porte da Empresa", ("Pequeno", "Médio", "Grande"))
        cidade = st.text_input("Município")
        corede = st.text_input("COREDE")
        qtd_empregos = st.text_input("Quantidade de Empregos", placeholder="1.000")
    with c4:
        pontos_fundopem = st.text_input("Pontos FUNDOPEM (3.1)")
        set_estrategicos = st.text_input("Pontos Set. Estratégicos")
        intensidade_tec = st.text_input("Pontos Intensidade Tecnológica")
        perc_integrar = st.text_input("Percentual INTEGRAR (3.2)", placeholder="34,55")
        pontos_idese = st.text_input("Pontos IDESE")
        pontos_setor = st.text_input("Pontos Setor Industrial")

    st.subheader("3. Valores e Prazos")
    c5, c6 = st.columns(2)
    with c5:
        valor_total = st.text_input("Valor Total do Projeto (2.1)")
        valor_apres_inicial = st.text_input("Valor Apresentado Inicialmente (2.3)")
        valor_inicial_aceito = st.text_input("Valor Inicialmente Aceito (2.3.1)")
        equips_24 = st.text_input("Equipamentos (2.4)")
        limite_max_liberado = st.text_input("Limite Máximo Liberado (4.1.2)")
        valor_liberado_fruicao = st.text_input("Valor Liberado p/ Fruição (4.1.2.1)")
    with c6:
        data_inicio = st.text_input("Início da Vigência (DD/MM/AAAA)", placeholder="01/08/2025")
        data_final_fruicao = st.text_input("Final da Fruição (DD/MM/AAAA)", placeholder="31/01/2032")
        mes_regularidade = st.text_input("Mês da Regularidade (ex: Julho/2025)")

    gerar = st.form_submit_button("🚀 Gerar Documento Word", type="primary")

# ─────────── PROCESSAMENTO ─────────── #
if gerar:
    with st.spinner("Gerando documento..."):
        mp: Dict[str, str] = {}

        # --- Valores (independentes) ---
        # 2.1 - Valor total do projeto
        if valor_total:
            fmt_total = br_currency(valor_total)
            mp["693.224,32 UIF/RS (seiscentos e noventa e três mil, duzentos e vinte e quatro inteiros, e trinta e dois centésimos de Unidades de Incentivo do FUNDOPEM/RS)"] = f"{fmt_total} UIF/RS ({num_extenso(fmt_total)} de Unidades de Incentivo do FUNDOPEM/RS)"
        
        # 2.3 - Valor apresentado inicialmente
        if valor_apres_inicial:
            fmt_apres = br_currency(valor_apres_inicial)
            mp["193.874,41 UIF/RS (cento e noventa e três mil, oitocentos e setenta e quatro inteiros, e quarenta e um centésimos de Unidades de Incentivo do FUNDOPEM/RS)"] = f"{fmt_apres} UIF/RS ({num_extenso(fmt_apres)} de Unidades de Incentivo do FUNDOPEM/RS)"
            mp["=g10"] = fmt_apres
        
        # 2.3.1 - Valor inicialmente aceito
        if valor_inicial_aceito:
            fmt_aceito = br_currency(valor_inicial_aceito)
            mp["159.123,22 UIF/RS (cento e cinquenta e nove mil, cento e vinte e três inteiros, e vinte e dois centésimos de Unidades de Incentivo do FUNDOPEM/RS)"] = f"{fmt_aceito} UIF/RS ({num_extenso(fmt_aceito)} de Unidades de Incentivo do FUNDOPEM/RS)"
            mp["=g11"] = fmt_aceito

        # 2.4 - Equipamentos
        if equips_24:
            fmt_equip = br_currency(equips_24)
            mp["Do valor estabelecido no item 2.3.1 desta Cláusula, o montante de 113.874,41 UIF/RS (cento e noventa e três mil, oitocentos e setenta e quatro inteiros, e quarenta e um centésimos de Unidades de Incentivo do FUNDOPEM/RS) contempla os investimentos realizados em equipamentos."] = (
                f"Do valor estabelecido no item 2.3.1 desta Cláusula, o montante de "
                f"{fmt_equip} UIF/RS ({num_extenso(fmt_equip)} de Unidades de Incentivo do FUNDOPEM/RS) contempla os investimentos realizados em equipamentos."
            )
        
        # 4.1.2 - Limite máximo liberado
        if limite_max_liberado:
            fmt_limite = br_currency(limite_max_liberado)
            mp["239.509,00 UIF/RS (duzentos e trinta e nove mil, e quinhentos e nove inteiros de Unidades de Incentivo do FUNDOPEM/RS)"] = f"{fmt_limite} UIF/RS ({num_extenso(fmt_limite)} de Unidades de Incentivo do FUNDOPEM/RS)"
            mp["=g8"] = fmt_limite
        
        # 4.1.2.1 - Valor liberado para fruição
        if valor_liberado_fruicao:
            fmt_fruicao = br_currency(valor_liberado_fruicao)
            mp["62.299,92 UIF/RS (sessenta e dois mil, duzentos e noventa e nove inteiros, e noventa e dois centésimos de Unidades de Incentivo do FUNDOPEM/RS)"] = f"{fmt_fruicao} UIF/RS ({num_extenso(fmt_fruicao)} de Unidades de Incentivo do FUNDOPEM/RS)"
            mp["=g21"] = fmt_fruicao

        # --- Pontos & percentual ---
        if pontos_fundopem:
            mp["#pontos# (sessenta) pontos"] = f"{pontos_fundopem} ({num_extenso(pontos_fundopem, 'geral')}) pontos"
        if perc_integrar:
            perc_ext = num_extenso(perc_integrar, "percent") + " por cento"
            mp["#integrar# (trinta e quatro vírgula cinquenta e cinco por cento)"] = f"{perc_integrar}% ({perc_ext})"
            mp["#integrar#%"] = f"{perc_integrar}%"
            mp["#integrar#"] = perc_integrar
            
        # --- Quantidade de empregos ---
        if qtd_empregos:
            mp["#emp#"] = str(qtd_empregos)
            mp["(duzentos e noventa e sete)"] = f"({num_extenso(qtd_empregos, 'geral')})"
            mp["(duzentos e noventa e quatro)"] = f"({num_extenso(qtd_empregos, 'geral')})"

        # --- Simples ---
        mp.update({
            "#xx/aaaa#": termo_num,
            "#EMPRESA#": empresa_nome.upper(),
            "#ENDERECO#": empresa_endereco,
            "#XX.XXX.XXX/0001-XX#": format_cnpj(empresa_cnpj),
            "#REPRESENTANTE#": representante_nome,
            "#CPF#": format_cpf(representante_cpf),
            "#cpf#": format_cpf(representante_cpf),
            "#proa#": proa_num,
            "#cidade#": cidade,
            "#corede#": corede,
            "#MUNICIPIO_E_COREDE#": f"{cidade}/RS (COREDE: {corede})" if cidade and corede else "",
            "#idese#": pontos_idese,
            "#setint#": pontos_setor,
            "#set#": set_estrategicos,
            "#it#": intensidade_tec,
            "#porte#": empresa_porte,
            "#cgcte#": empresa_cgcte,
            "#pontos#": pontos_fundopem,
        })

        mp["CPF XXX.XXX.XXX-XX"] = f"CPF {format_cpf(representante_cpf)}"
        mp.pop("#CPF#", None)
        mp.pop("#cpf#", None)

        # --- Datas ---
        # Usando formato com ponto para as datas do parecer
        dt_inicio = to_date(data_inicio)
        dt_final = to_date(data_final_fruicao)
        dt_parecer = to_date(parecer_data)
        dt_doe = to_date(doe_data)
        dt_proa = to_date(proa_data)
        
        # Data de início em formato com ponto (01.08.2025)
        mp["#inicio#"] = data_formato_ponto(dt_inicio)
        
        # Data final em formato Mês/Ano
        mp["#final#"] = data_texto(dt_final, mes_ano=True)
        mp["#final2#"] = data_texto(dt_final, mes_ano=True).replace(" ", "")
        mp["#regularidade#"] = mes_regularidade

        # Parecer com datas em formato ponto
        if parecer_num and dt_parecer and dt_doe:
            mp["Parecer nº xxx/aaaa, de dd.mm.aaaa (DOE de dd.mm.aaaa)"] = (
                f"Parecer nº {parecer_num}, de {data_formato_ponto(dt_parecer)} "
                f"(DOE de {data_formato_ponto(dt_doe)})"
            )
        
        # Corrigir a data do processo PROA - procurando pelo texto exato do template
        if proa_num and dt_proa:
            # O template tem o texto completo incluindo a vírgula dupla
            mp["e na documentação que instrui o processo administrativo nº #proa#, de 03 de setembro de 2024, que passam a fazer parte integrante deste instrumento."] = (
                f"e na documentação que instrui o processo administrativo nº {proa_num}, de {data_formato_ponto(dt_proa)}, que passam a fazer parte integrante deste instrumento."
            )

        # --- DEBUG: Mostrar substituições ---
        with st.expander("🔍 Debug - Substituições a serem realizadas"):
            st.write("### Principais substituições:")
            
            # Valores
            st.write("**Valores:**")
            for k, v in mp.items():
                if "UIF/RS" in k and len(k) > 50:
                    st.code(f"{k[:80]}... → {v}")
            
            # Placeholders simples
            st.write("\n**Placeholders simples:**")
            placeholders = [k for k in mp.keys() if "#" in k or "xxx" in k or "XXX" in k]
            for k in sorted(placeholders):
                st.code(f"{k} → {mp[k]}")
            
            # Datas
            st.write("\n**Datas e textos com datas:**")
            for k, v in mp.items():
                if "dd.mm.aaaa" in k or "de 03 de setembro" in k or "Parecer" in k:
                    st.code(f"{k[:100]}... → {v}")
            
            # Total de substituições
            st.info(f"Total de substituições a serem realizadas: {len(mp)}")

        # --- Geração DOCX ---
        doc = Document(str(TEMPLATE_FILE))
        docx_replace(doc, mp)
        buf = io.BytesIO()
        doc.save(buf)

        nome = f"Termo_Ajuste_{(termo_num or 'novo').replace('/', '-')}_{(empresa_nome or 'empresa').replace(' ', '_')}.docx"
        st.success("🎉 Documento gerado com sucesso!")
        st.download_button("⬇️ Baixar documento Word", buf.getvalue(), nome,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                          use_container_width=True)